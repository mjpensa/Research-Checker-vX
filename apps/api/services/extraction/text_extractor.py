import aiofiles
from pathlib import Path
from typing import Optional
import logging
import json
from dataclasses import dataclass

logger = logging.getLogger(__name__)

@dataclass
class ExtractedDocument:
    text: str
    metadata: dict
    page_count: Optional[int] = None
    word_count: Optional[int] = None

class TextExtractor:
    """Extract text from various document formats"""

    async def extract(self, file_path: str, mime_type: str) -> ExtractedDocument:
        """Extract text based on file type"""

        path = Path(file_path)

        if mime_type == 'application/pdf':
            return await self._extract_pdf(path)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return await self._extract_docx(path)
        elif mime_type in ['text/markdown', 'text/plain']:
            return await self._extract_text(path)
        elif mime_type == 'application/json':
            return await self._extract_json(path)
        else:
            raise ValueError(f"Unsupported mime type: {mime_type}")

    async def _extract_pdf(self, path: Path) -> ExtractedDocument:
        """Extract text from PDF using pdfplumber"""
        try:
            import pdfplumber

            text_parts = []
            page_count = 0

            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)

                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)

            return ExtractedDocument(
                text=full_text,
                metadata={'source_format': 'pdf'},
                page_count=page_count,
                word_count=len(full_text.split())
            )

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

    async def _extract_docx(self, path: Path) -> ExtractedDocument:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document

            doc = Document(path)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_texts.append(row_text)

            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n=== TABLES ===\n\n" + "\n".join(table_texts)

            return ExtractedDocument(
                text=full_text,
                metadata={'source_format': 'docx'},
                page_count=len(doc.sections),
                word_count=len(full_text.split())
            )

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise

    async def _extract_text(self, path: Path) -> ExtractedDocument:
        """Extract text from plain text or markdown"""
        try:
            async with aiofiles.open(path, 'r', encoding='utf-8') as f:
                text = await f.read()

            return ExtractedDocument(
                text=text,
                metadata={'source_format': path.suffix[1:]},
                word_count=len(text.split())
            )

        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            raise

    async def _extract_json(self, path: Path) -> ExtractedDocument:
        """Extract text from JSON (assumes specific structure)"""
        try:
            async with aiofiles.open(path, 'r', encoding='utf-8') as f:
                content = await f.read()
                data = json.loads(content)

            # Extract text based on common structures
            if isinstance(data, dict):
                if 'text' in data:
                    text = data['text']
                elif 'content' in data:
                    text = data['content']
                elif 'messages' in data:
                    # Chat format
                    text = "\n\n".join([
                        f"{msg.get('role', 'unknown')}: {msg.get('content', '')}"
                        for msg in data['messages']
                    ])
                else:
                    # Generic JSON - extract all string values
                    text = self._extract_strings_from_dict(data)
            elif isinstance(data, list):
                text = "\n\n".join(str(item) for item in data)
            else:
                text = str(data)

            return ExtractedDocument(
                text=text,
                metadata={'source_format': 'json', 'structure': type(data).__name__},
                word_count=len(text.split())
            )

        except Exception as e:
            logger.error(f"JSON extraction error: {e}")
            raise

    def _extract_strings_from_dict(self, obj, depth=0, max_depth=5):
        """Recursively extract string values from nested dict/list"""
        if depth > max_depth:
            return ""

        parts = []

        if isinstance(obj, dict):
            for value in obj.values():
                if isinstance(value, str):
                    parts.append(value)
                elif isinstance(value, (dict, list)):
                    parts.append(self._extract_strings_from_dict(value, depth + 1))
        elif isinstance(obj, list):
            for item in obj:
                if isinstance(item, str):
                    parts.append(item)
                elif isinstance(item, (dict, list)):
                    parts.append(self._extract_strings_from_dict(item, depth + 1))

        return " ".join(parts)

text_extractor = TextExtractor()
